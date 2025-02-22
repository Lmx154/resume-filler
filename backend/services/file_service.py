import os
import PyPDF2
import docx
from pathlib import Path
from typing import Dict
import re

class FileService:
    def read_file_content(self, file_path: str) -> str:
        """Read content from PDF, DOCX, or TXT files"""
        ext = Path(file_path).suffix.lower()
        if ext == '.pdf':
            return self._read_pdf(file_path)
        elif ext == '.docx':
            return self._read_docx(file_path)
        elif ext == '.txt':
            return self._read_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    # New helper to read PDF from a file-like object (bytes)
    def _read_pdf_from_bytes(self, file_obj) -> str:
        reader = PyPDF2.PdfReader(file_obj)
        return ' '.join(page.extract_text() for page in reader.pages if page.extract_text())

    # New helper to read DOCX from a file-like object (bytes)
    def _read_docx_from_bytes(self, file_obj) -> str:
        doc = docx.Document(file_obj)
        return ' '.join(paragraph.text for paragraph in doc.paragraphs)

    def _read_pdf(self, file_path: str) -> str:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            return ' '.join(page.extract_text() for page in reader.pages if page.extract_text())

    def _read_docx(self, file_path: str) -> str:
        doc = docx.Document(file_path)
        return ' '.join(paragraph.text for paragraph in doc.paragraphs)

    def _read_txt(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

    def parse_resume(self, content: str) -> Dict:
        """Parse resume content into a structured format"""
        cleaned_content = self._clean_text(content)
        sections = self._identify_sections(cleaned_content)
        return {
            "parsed_sections": sections,
            "summary": self._create_summary(sections),
            "metadata": self._extract_metadata(cleaned_content)
        }

    def _clean_text(self, text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[\r\n]+', '\n', text)
        return text.strip()

    def _identify_sections(self, text: str) -> Dict:
        common_headers = [
            "experience", "education", "skills", "projects",
            "summary", "objective", "certifications"
        ]
        sections = {}
        current_section = "other"
        current_content = []
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
            line_lower = line.lower()
            found_header = False
            for header in common_headers:
                if header in line_lower and len(line) < 50:
                    if current_content:
                        sections[current_section] = '\n'.join(current_content)
                    current_section = header
                    current_content = []
                    found_header = True
                    break
            if not found_header:
                current_content.append(line)
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        return sections

    def _create_summary(self, sections: Dict) -> str:
        summary_parts = []
        if "summary" in sections:
            summary_parts.append(sections["summary"])
        if "experience" in sections:
            exp_sentences = re.split(r'[.!?]+', sections["experience"])
            summary_parts.append(' '.join(exp_sentences[:2]))
        if "skills" in sections:
            summary_parts.append(f"Key skills include: {sections['skills']}")
        return ' '.join(summary_parts)

    def _extract_metadata(self, text: str) -> Dict:
        words = text.split()
        sentences = re.split(r'[.!?]+', text)
        return {
            "word_count": len(words),
            "sentence_count": len(sentences),
            "estimated_read_time": len(words) // 200
        }

file_service = FileService()